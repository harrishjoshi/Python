#! python3

# custom_invitations.py - Generates a word document containing
# custom invitations, one per page, using names from guests.txt.

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Create a new Word document
doc = Document()

# Create style for invitation body text
body_style = doc.styles.add_style("InviteBody", 1)
body_style.font.name = "Times New Roman"
body_style.font.size = Pt(18)
body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Create style for guest names
name_style = doc.styles.add_style("InviteName", 1)
name_style.font.name = "Times New Roman"
name_style.font.size = Pt(28)
name_style.font.bold = True
name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Open the guests.txt file and read each name into a list
with open("files/guests.txt", "r") as file:
    guests = [line.strip() for line in file if line.strip()]

# Loop through each guest name
for guest in guests:
    # Add opening line of the invitation
    doc.add_paragraph(
        "It would be a pleasure to have the company of", style="InviteBody"
    )

    # Add the guest's name
    doc.add_paragraph(guest, style="InviteName")

    # Add the event details
    doc.add_paragraph(
        "at <Location-Name> the Evening of <Date>\nat <Time>", style="InviteBody"
    )

    # Add closing line
    doc.add_paragraph("Looking forward to seeing you,\n\nHost Name", style="InviteBody")

    # Insert a page break after each invitation
    doc.add_page_break()

# Save the completed Word document
doc.save("files/custom_invitations.docx")
